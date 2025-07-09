import markdown
from docx import Document
import re
import sys
import os

def md_to_docx(md_file, docx_file):
    # 读取Markdown文件
    with open(md_file, 'r', encoding='utf-8') as f:
        md_content = f.read()
    
    # 转换Markdown为HTML
    html = markdown.markdown(md_content, extensions=['tables', 'fenced_code'])
    
    # 创建Word文档
    doc = Document()
    
    # 添加标题
    title = os.path.splitext(os.path.basename(md_file))[0]
    doc.add_heading('基于神经网络预测的飞控系统闭环控制技术文档', 0)
    
    # 简单处理HTML并添加到Word文档
    # 注意：这是一个非常基本的转换，不能处理所有HTML元素
    paragraphs = html.split('<p>')
    for p in paragraphs:
        if p.strip():
            clean_p = re.sub('<[^<]+?>', '', p)
            doc.add_paragraph(clean_p)
    
    # 保存Word文档
    doc.save(docx_file)
    print(f"已将 {md_file} 转换为 {docx_file}")

if __name__ == "__main__":
    if len(sys.argv) > 1:
        md_file = sys.argv[1]
    else:
        md_file = "1.md"  # 默认文件名
    
    docx_file = "神经网络飞控系统文档.docx"
    md_to_docx(md_file, docx_file) 