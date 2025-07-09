import markdown
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re
import os
from bs4 import BeautifulSoup

def md_to_docx_local(md_file, docx_file):
    """将Markdown文件转换为Word文档，专注于表格渲染"""
    # 读取Markdown文件
    with open(md_file, 'r', encoding='utf-8') as f:
        md_content = f.read()
    
    # 转换Markdown为HTML，启用表格扩展
    html = markdown.markdown(md_content, extensions=['tables', 'fenced_code'])
    
    # 解析HTML
    soup = BeautifulSoup(html, 'html.parser')
    
    # 创建Word文档
    doc = Document()
    
    # 添加标题
    doc.add_heading('基于神经网络预测的飞控系统闭环控制技术文档', 0)
    
    # 处理段落、标题和表格
    for element in soup.find_all(['p', 'h1', 'h2', 'h3', 'h4', 'h5', 'h6', 'table', 'pre']):
        if element.name.startswith('h'):
            level = int(element.name[1])
            doc.add_heading(element.text, level)
        elif element.name == 'p':
            # 检查是否包含mermaid标记
            if 'mermaid' in element.text:
                p = doc.add_paragraph("【此处应有流程图 - 请参考原Markdown文件】")
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                # 添加一个空行
                doc.add_paragraph("")
            else:
                doc.add_paragraph(element.text)
        elif element.name == 'pre':
            # 处理代码块
            code_text = element.text
            if 'mermaid' in code_text:
                p = doc.add_paragraph("【此处应有流程图 - 请参考原Markdown文件】")
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                p = doc.add_paragraph(code_text)
                p.style = 'Code'
        elif element.name == 'table':
            # 处理表格
            rows = element.find_all('tr')
            if rows:
                # 获取最大列数
                max_cols = max(len(row.find_all(['td', 'th'])) for row in rows)
                
                # 创建表格
                table = doc.add_table(rows=len(rows), cols=max_cols)
                table.style = 'Table Grid'
                
                # 填充表格内容
                for i, row in enumerate(rows):
                    cells = row.find_all(['td', 'th'])
                    for j, cell in enumerate(cells):
                        if j < max_cols:  # 确保不超出列数
                            table.cell(i, j).text = cell.text.strip()
                
                # 表格后添加空行
                doc.add_paragraph()
    
    # 保存Word文档
    doc.save(docx_file)
    print(f"已将 {md_file} 转换为 {docx_file}，表格已正确渲染")

if __name__ == "__main__":
    md_file = "1.md"
    docx_file = "神经网络飞控系统文档_表格版.docx"
    md_to_docx_local(md_file, docx_file) 