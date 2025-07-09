import markdown
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re
import os
import base64
import requests
from bs4 import BeautifulSoup
import io
import tempfile
from PIL import Image

def convert_mermaid_to_image(mermaid_code):
    """将Mermaid代码转换为图片URL"""
    try:
        # 使用Mermaid Live Editor API
        graphbytes = mermaid_code.encode("utf8")
        base64_bytes = base64.b64encode(graphbytes)
        base64_string = base64_bytes.decode("ascii")
        
        url = f"https://mermaid.ink/img/{base64_string}"
        response = requests.get(url)
        
        if response.status_code == 200:
            # 保存为临时文件
            img = Image.open(io.BytesIO(response.content))
            temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.png')
            img.save(temp_file.name)
            return temp_file.name
        else:
            print(f"无法转换Mermaid图表，状态码: {response.status_code}")
            return None
    except Exception as e:
        print(f"转换Mermaid图表时出错: {e}")
        return None

def md_to_docx_advanced(md_file, docx_file):
    # 读取Markdown文件
    with open(md_file, 'r', encoding='utf-8') as f:
        md_content = f.read()
    
    # 提取所有mermaid代码块
    mermaid_blocks = re.findall(r'```mermaid(.*?)```', md_content, re.DOTALL)
    mermaid_images = []
    
    # 转换mermaid代码块为图片
    for i, block in enumerate(mermaid_blocks):
        img_path = convert_mermaid_to_image(block)
        if img_path:
            mermaid_images.append(img_path)
            # 替换mermaid代码块为占位符
            md_content = md_content.replace(f"```mermaid{block}```", f"[MERMAID_IMAGE_{i}]")
    
    # 转换Markdown为HTML
    html = markdown.markdown(md_content, extensions=['tables', 'fenced_code'])
    
    # 解析HTML
    soup = BeautifulSoup(html, 'html.parser')
    
    # 创建Word文档
    doc = Document()
    
    # 添加标题
    doc.add_heading('基于神经网络预测的飞控系统闭环控制技术文档', 0)
    
    # 处理段落和标题
    for element in soup.find_all(['p', 'h1', 'h2', 'h3', 'h4', 'h5', 'h6', 'table']):
        if element.name.startswith('h'):
            level = int(element.name[1])
            doc.add_heading(element.text, level)
        elif element.name == 'p':
            # 检查是否为mermaid图片占位符
            match = re.search(r'\[MERMAID_IMAGE_(\d+)\]', element.text)
            if match:
                img_index = int(match.group(1))
                if img_index < len(mermaid_images):
                    doc.add_picture(mermaid_images[img_index], width=Inches(6))
                    paragraph = doc.paragraphs[-1]
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                doc.add_paragraph(element.text)
        elif element.name == 'table':
            # 处理表格
            rows = element.find_all('tr')
            if rows:
                table = doc.add_table(rows=len(rows), cols=len(rows[0].find_all(['td', 'th'])))
                table.style = 'Table Grid'
                
                for i, row in enumerate(rows):
                    cells = row.find_all(['td', 'th'])
                    for j, cell in enumerate(cells):
                        try:
                            table.cell(i, j).text = cell.text.strip()
                        except IndexError:
                            pass  # 处理表格列数不一致的情况
    
    # 保存Word文档
    doc.save(docx_file)
    print(f"已将 {md_file} 转换为 {docx_file}")
    
    # 清理临时文件
    for img_path in mermaid_images:
        try:
            os.remove(img_path)
        except:
            pass

if __name__ == "__main__":
    md_file = "1.md"
    docx_file = "神经网络飞控系统文档_高级版.docx"
    md_to_docx_advanced(md_file, docx_file) 